"""
创建测试用的Word文档
用于测试文档解析功能
"""

from docx import Document
from docx.shared import Inches
import os

def create_test_word_document():
    """创建测试用的Word文档"""
    
    # 创建新文档
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('某某设计院 - IO点表清单', 0)
    
    # 添加说明段落
    doc.add_paragraph('本文档包含项目的IO点位信息，用于测试文档解析功能。')
    
    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '仪表位号'
    header_cells[1].text = '检测点名称'
    header_cells[2].text = '信号类型'
    header_cells[3].text = '量程'
    header_cells[4].text = '单位'
    header_cells[5].text = '备注'
    
    # 添加测试数据
    test_data = [
        ['PT-1101', '进站压力检测', '4-20mA', '0-1.6', 'MPa', '压力变送器'],
        ['TT-1102', '进站温度检测', '4-20mA', '0-100', '℃', '温度变送器'],
        ['FT-1103', '主管流量检测', '4-20mA', '0-1000', 'm³/h', '流量计'],
        ['LT-1104', '储罐液位检测', '4-20mA', '0-5', 'm', '液位计'],
        ['UA-1201', '进站压力高报警', '开关量', '断开/闭合', '', '压力开关'],
        ['UA-1202', '出站故障报警', '干接点', '正常/故障', '', '故障指示'],
        ['XS-1203', '阀门位置反馈', '开关量', '开/关', '', '位置开关'],
        ['XO-1301', '进站紧急切断阀', '0/24VDC', '开/关', '', '电动阀'],
        ['XO-1302', '出站控制阀', '继电器', '开/关', '', '电磁阀'],
        ['YO-1303', '加热器控制', '220VAC', '启动/停止', '', '加热器'],
        ['AO-1401', '调节阀控制', '4-20mA输出', '0-100', '%', '调节阀'],
        ['AO-1402', '变频器设定', '0-10V输出', '0-50', 'Hz', '变频器']
    ]
    
    # 添加数据行
    for data_row in test_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(data_row):
            row_cells[i].text = cell_data
    
    # 保存文档
    output_dir = 'test_documents'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    file_path = os.path.join(output_dir, '测试IO点表清单.docx')
    doc.save(file_path)
    
    print(f"测试Word文档已创建: {file_path}")
    return file_path

if __name__ == "__main__":
    create_test_word_document()
