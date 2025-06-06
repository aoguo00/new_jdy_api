"""
Word文档解析器
用于解析设计院提供的Word格式的IO点表文档
"""

import logging
from typing import List, Dict, Any, Optional
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base_parser import BaseDocumentParser, ParseError, UnsupportedFormatError
from .smart_header_detector import SmartHeaderDetector

logger = logging.getLogger(__name__)


class WordDocumentParser(BaseDocumentParser):
    """Word文档解析器"""
    
    def __init__(self):
        """初始化Word文档解析器"""
        super().__init__()
        self.supported_extensions = ['.docx']
        self.header_detector = SmartHeaderDetector()
        
    def parse_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析Word文档并提取点位信息
        
        Args:
            file_path (str): Word文档文件路径
            
        Returns:
            List[Dict[str, Any]]: 提取的点位信息列表
        """
        logger.info(f"开始解析Word文档: {file_path}")
        
        # 验证文件
        if not self.validate_file(file_path):
            raise UnsupportedFormatError(f"文件验证失败: {file_path}")
            
        try:
            # 打开Word文档
            document = Document(file_path)
            logger.info(f"成功打开Word文档，包含 {len(document.tables)} 个表格")
            
            # 提取所有点位信息
            all_points = []
            
            # 遍历文档中的所有表格
            for table_index, table in enumerate(document.tables):
                logger.info(f"正在处理第 {table_index + 1} 个表格")
                
                try:
                    points = self._parse_table(table, table_index)
                    if points:
                        all_points.extend(points)
                        logger.info(f"从第 {table_index + 1} 个表格中提取到 {len(points)} 个点位")
                    else:
                        logger.info(f"第 {table_index + 1} 个表格未识别到有效点位信息")
                        
                except Exception as e:
                    logger.warning(f"解析第 {table_index + 1} 个表格时出错: {e}")
                    continue
                    
            # 标准化点位数据
            standardized_points = []
            for point in all_points:
                try:
                    standard_point = self.standardize_point_data(point)
                    standardized_points.append(standard_point)
                except Exception as e:
                    logger.warning(f"标准化点位数据时出错: {e}")
                    continue
                    
            logger.info(f"Word文档解析完成，共提取到 {len(standardized_points)} 个有效点位")
            return standardized_points
            
        except Exception as e:
            logger.error(f"解析Word文档失败: {e}")
            raise ParseError(f"解析Word文档失败: {str(e)}")
            
    def _parse_table(self, table: Table, table_index: int) -> List[Dict[str, Any]]:
        """
        解析单个表格
        
        Args:
            table (Table): Word表格对象
            table_index (int): 表格索引
            
        Returns:
            List[Dict[str, Any]]: 从表格中提取的点位信息
        """
        if len(table.rows) < 2:
            logger.debug(f"表格 {table_index + 1} 行数不足，跳过")
            return []
            
        # 使用智能表头检测器识别表头
        header_mapping = self._identify_table_headers_smart(table)
        if not header_mapping:
            logger.debug(f"表格 {table_index + 1} 未识别到有效表头，跳过")
            return []

        logger.info(f"表格 {table_index + 1} 识别到的列映射: {header_mapping}")
        
        # 找到表头行位置
        header_row_index = self._find_header_row(table)
        if header_row_index is None:
            header_row_index = 0

        # 从表头行的下一行开始提取数据
        start_row = header_row_index + 1
        print(f"从第 {start_row} 行开始提取数据...")

        # 提取数据行
        points = []
        for row_index in range(start_row, len(table.rows)):
            try:
                row = table.rows[row_index]
                point_data = self._extract_row_data(row, header_mapping, row_index)
                # 如果 _extract_row_data 返回 None，说明是被过滤的行
                if point_data is None:
                    continue
                if point_data and self._is_valid_point_data(point_data):
                    points.append(point_data)

            except Exception as e:
                logger.warning(f"提取表格 {table_index + 1} 第 {row_index + 1} 行数据时出错: {e}")
                continue
                
        return points

    def _identify_table_headers_smart(self, table: Table) -> Dict[str, int]:
        """
        使用智能检测器识别表格表头

        Args:
            table (Table): Word表格对象

        Returns:
            Dict[str, int]: 字段名到列索引的映射
        """
        # 找到表头行
        header_row_index = self._find_header_row(table)
        if header_row_index is None:
            print("❌ 未找到有效的表头行")
            return {}

        header_row = table.rows[header_row_index]

        # 提取表头文本
        header_texts = []
        for cell in header_row.cells:
            cell_text = self._get_cell_text(cell).strip()
            header_texts.append(cell_text)

        # 使用智能检测器进行表头识别
        header_mapping = self.header_detector.detect_headers(header_texts)

        return header_mapping

    def _identify_table_headers(self, table: Table) -> Dict[str, int]:
        """
        识别表格表头并建立列映射

        Args:
            table (Table): Word表格对象

        Returns:
            Dict[str, int]: 列名到列索引的映射
        """
        if not table.rows:
            return {}

        # 先找到真正的表头行
        header_row_index = self._find_header_row(table)
        if header_row_index is None:
            print("❌ 未找到有效的表头行")
            return {}

        header_row = table.rows[header_row_index]
        header_mapping = {}
        
        # 定义关键字映射 - 扩展更多可能的表头
        keyword_mappings = {
            'instrument_tag': ['位号', 'tag', '标号', '仪表位号', '设备位号', '编号', '序号', 'NO', 'No', '点位号'],
            'description': ['名称', '描述', 'description', 'name', '检测点', '测点', '点位名称', '变量名', '功能描述', '说明'],
            'signal_range': ['信号范围'],  # 信号范围 (如 4~20mA) - 精确匹配
            'data_range': ['数据范围', '量程', 'range', '范围', '数据'],  # 数据范围 (如 0-6, -20~80)
            'signal_type': ['信号类型', '类型', 'type', '信号'],  # 信号类型 (如 AI, DI) - 信号优先匹配到这里
            'units': ['单位', 'unit', '量纲'],  # 单位 (如 MPa, ℃)
            'power_supply': ['现场仪表供电', '供电', 'power', '电源', '仪表供电'],  # 现场仪表供电
            'isolation': ['隔离', 'isolation', '隔离器'],  # 隔离
            'io_type': ['io', 'IO', 'I/O', '通道类型', '通道', '模块'],
            'range_low': ['下限', 'low', 'min', '最小值', '量程下限'],
            'range_high': ['上限', 'high', 'max', '最大值', '量程上限'],
            'location': ['位置', '安装位置', 'location', '安装地点'],
            'remarks': ['备注', '说明', 'remarks', 'note', '注释']
        }
        
        # 先打印所有表头内容用于调试
        print(f"\n=== 表格表头调试信息 ===")
        header_texts = []
        for col_index, cell in enumerate(header_row.cells):
            cell_text = self._get_cell_text(cell).strip()
            header_texts.append(cell_text)
            print(f"列 {col_index}: '{cell_text}'")
        print(f"表头总数: {len(header_texts)}")

        # 遍历表头单元格
        for col_index, cell in enumerate(header_row.cells):
            cell_text = self._get_cell_text(cell).strip()
            if not cell_text:
                continue

            # 匹配关键字 - 按照精确度排序，优先匹配更精确的字段
            matched = False
            best_match = None
            best_score = 0

            for field_name, keywords in keyword_mappings.items():
                for keyword in keywords:
                    if keyword in cell_text:
                        # 计算匹配得分：完全匹配得分更高，长关键字得分更高
                        if cell_text == keyword:
                            score = len(keyword) * 2  # 完全匹配得分翻倍
                        else:
                            score = len(keyword)  # 部分匹配

                        if score > best_score:
                            best_score = score
                            best_match = (field_name, keyword)

            if best_match:
                field_name, keyword = best_match
                header_mapping[field_name] = col_index
                print(f"✅ 识别到列 '{cell_text}' -> {field_name} (索引: {col_index})")
                matched = True

            if not matched:
                print(f"❌ 未识别列 '{cell_text}' (索引: {col_index})")
                    
        # 检查是否识别到足够的关键列
        print(f"识别到的字段映射: {header_mapping}")

        # 如果没有识别到任何字段，尝试使用默认映射
        if not header_mapping:
            print("❌ 没有识别到任何有效字段，尝试使用默认映射")
            if len(header_texts) >= 2:
                header_mapping = {
                    'instrument_tag': 0,  # 第一列作为位号
                    'description': 1      # 第二列作为描述
                }
                print(f"⚠️ 使用默认映射: 第1列=位号, 第2列=描述")
            else:
                print("❌ 表格列数不足，无法使用默认映射")
                return {}

        print(f"✅ 最终字段映射: {header_mapping}")
        return header_mapping

    def _find_header_row(self, table: Table) -> Optional[int]:
        """
        查找真正的表头行

        Args:
            table (Table): Word表格对象

        Returns:
            Optional[int]: 表头行索引，如果未找到则返回None
        """
        print(f"\n=== 查找表头行 ===")

        # 检查前10行，寻找真正的表头
        for row_index in range(min(10, len(table.rows))):
            row = table.rows[row_index]
            row_texts = [self._get_cell_text(cell).strip() for cell in row.cells]
            row_text = ' '.join(row_texts)

            print(f"第 {row_index} 行: {row_texts[:8]}...")  # 显示前8列

            # 检查是否包含表头关键字
            header_keywords = ['仪表位号', '位号', '检测点名称', '名称', '信号类型', '信号范围', '数据范围', '信号', '通道类型', '量程', '单位', '现场仪表供电', '供电', '隔离']
            matched_keywords = [kw for kw in header_keywords if kw in row_text]

            if len(matched_keywords) >= 2:  # 至少匹配2个关键字
                print(f"✅ 识别到表头行: 第 {row_index} 行，匹配关键字: {matched_keywords}")
                return row_index

        # 如果没有找到明确的表头，尝试找包含"位号"的行
        for row_index in range(min(10, len(table.rows))):
            row = table.rows[row_index]
            row_text = ' '.join([self._get_cell_text(cell).strip() for cell in row.cells])
            if '位号' in row_text:
                print(f"⚠️ 找到包含'位号'的行: 第 {row_index} 行")
                return row_index

        print("❌ 未找到明确表头，使用第一行")
        return 0
        
    def _extract_row_data(self, row, header_mapping: Dict[str, int], row_index: int) -> Dict[str, Any]:
        """
        从表格行中提取数据

        Args:
            row: 表格行对象
            header_mapping (Dict[str, int]): 列映射
            row_index (int): 行索引

        Returns:
            Dict[str, Any]: 提取的行数据
        """
        row_data = {'_row_index': row_index}

        for field_name, col_index in header_mapping.items():
            try:
                if col_index < len(row.cells):
                    cell_text = self._get_cell_text(row.cells[col_index]).strip()
                    row_data[field_name] = cell_text
                else:
                    row_data[field_name] = ''

            except Exception as e:
                logger.warning(f"提取第 {row_index} 行第 {col_index} 列数据时出错: {e}")
                row_data[field_name] = ''

        # 在这里进行早期过滤，检查是否为分组标题行
        tag = row_data.get('instrument_tag', '').strip()
        desc = row_data.get('description', '').strip()

        # 检查是否为分组标题行（系统名称）
        system_names = ['BPCS', 'ESD', 'RS485', 'DCS', 'SIS', 'F&G', 'FIRE', 'GAS']
        if tag.upper() in system_names and not desc:
            print(f"  -> 跳过分组标题行: {tag}")
            return None

        # 检查是否为表格设计标题行
        design_keywords = ['设计', '审核', '校对', '批准', '设 计', '审 核']
        if any(keyword in tag or keyword in desc for keyword in design_keywords):
            print(f"  -> 跳过设计标题行: {tag} | {desc}")
            return None

        return row_data
        
    def _get_cell_text(self, cell) -> str:
        """
        获取单元格文本内容
        
        Args:
            cell: 单元格对象
            
        Returns:
            str: 单元格文本
        """
        try:
            # 获取单元格中所有段落的文本
            text_parts = []
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
                    
            return ' '.join(text_parts)
            
        except Exception as e:
            logger.warning(f"获取单元格文本时出错: {e}")
            return ''
            
    def _is_valid_point_data(self, point_data: Dict[str, Any]) -> bool:
        """
        验证点位数据是否有效

        Args:
            point_data (Dict[str, Any]): 点位数据

        Returns:
            bool: 数据是否有效
        """
        # 检查是否有基本的标识信息
        has_tag = bool(point_data.get('instrument_tag', '').strip())
        has_description = bool(point_data.get('description', '').strip())

        tag = point_data.get('instrument_tag', '').strip()
        desc = point_data.get('description', '').strip()

        # 检查是否为分组标题行（系统名称）
        system_names = ['BPCS', 'ESD', 'RS485', 'DCS', 'SIS', 'F&G', 'FIRE', 'GAS']
        if tag.upper() in system_names and not desc:
            print(f"  -> 跳过分组标题行: {tag}")
            return False

        # 检查是否为表格设计标题行
        design_keywords = ['设计', '审核', '校对', '批准', '设 计', '审 核']
        if any(keyword in tag or keyword in desc for keyword in design_keywords):
            print(f"  -> 跳过设计标题行: {tag} | {desc}")
            return False

        # 至少需要有位号或描述之一
        if not (has_tag or has_description):
            return False

        # 检查是否为表头行或空行
        tag_lower = tag.lower()
        desc_lower = desc.lower()

        # 过滤掉明显的表头行
        header_keywords = ['位号', 'tag', '名称', 'name', '序号', 'no', '编号']
        if any(keyword in tag_lower or keyword in desc_lower for keyword in header_keywords):
            return False

        return True
